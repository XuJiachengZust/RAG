from typing import Dict, List, Optional, Any
from pathlib import Path
import logging
import tempfile
import os

# 文档转换库
try:
    import pypandoc
except ImportError:
    pypandoc = None

try:
    from docx import Document
except ImportError:
    Document = None

try:
    import PyPDF2
    from PyPDF2 import PdfReader
except ImportError:
    PyPDF2 = None
    PdfReader = None

logger = logging.getLogger(__name__)

class DocumentConverter:
    """
    文档转换器 - 将Word和PDF文档转换为Markdown格式
    高内聚的文档转换逻辑，专注于格式转换
    """
    
    def __init__(self, config: Optional[Dict[str, Any]] = None):
        """
        初始化文档转换器
        
        Args:
            config: 配置字典
        """
        self.config = config or {}
        self.temp_dir = tempfile.mkdtemp(prefix="doc_converter_")
        
        # 检查依赖
        self._check_dependencies()
    
    def _check_dependencies(self) -> None:
        """
        检查必要的依赖库
        """
        missing_deps = []
        
        if pypandoc is None:
            missing_deps.append("pypandoc")
        
        if Document is None:
            missing_deps.append("python-docx")
            
        if PyPDF2 is None:
            missing_deps.append("PyPDF2")
            
        if missing_deps:
            logger.warning(f"缺少依赖库: {', '.join(missing_deps)}. 某些转换功能可能不可用。")
    
    def convert_to_markdown(self, file_path: str) -> str:
        """
        将文档转换为Markdown格式
        
        Args:
            file_path: 文档文件路径
            
        Returns:
            转换后的Markdown内容
            
        Raises:
            ValueError: 不支持的文件格式
            RuntimeError: 转换失败
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")
        
        file_extension = file_path.suffix.lower()
        
        try:
            if file_extension == '.docx':
                return self._convert_docx_to_markdown(file_path)
            elif file_extension == '.pdf':
                return self._convert_pdf_to_markdown(file_path)
            else:
                raise ValueError(f"不支持的文件格式: {file_extension}")
                
        except Exception as e:
            logger.error(f"文档转换失败 {file_path}: {str(e)}")
            raise RuntimeError(f"转换失败: {str(e)}") from e
    
    def _convert_docx_to_markdown(self, file_path: Path) -> str:
        """
        将DOCX文档转换为Markdown
        
        Args:
            file_path: DOCX文件路径
            
        Returns:
            Markdown内容
        """
        if pypandoc is not None:
            # 使用pypandoc进行转换（推荐方式）
            try:
                markdown_content = pypandoc.convert_file(
                    str(file_path), 
                    'markdown',
                    format='docx',
                    extra_args=['--wrap=none', '--extract-media=' + self.temp_dir]
                )
                return markdown_content
            except Exception as e:
                logger.warning(f"pypandoc转换失败，使用备用方法: {str(e)}")
        
        # 备用方法：使用python-docx
        if Document is not None:
            return self._docx_to_markdown_fallback(file_path)
        
        raise RuntimeError("无可用的DOCX转换方法")
    
    def _docx_to_markdown_fallback(self, file_path: Path) -> str:
        """
        使用python-docx的备用转换方法
        
        Args:
            file_path: DOCX文件路径
            
        Returns:
            Markdown内容
        """
        doc = Document(file_path)
        markdown_lines = []
        
        # 处理文档中的所有元素（段落和表格）
        for element in doc.element.body:
            if element.tag.endswith('p'):  # 段落
                # 找到对应的段落对象
                for paragraph in doc.paragraphs:
                    if paragraph._element == element:
                        text = paragraph.text.strip()
                        if not text:
                            continue
                            
                        # 简单的样式转换
                        style_name = paragraph.style.name.lower()
                        
                        if 'heading 1' in style_name:
                            markdown_lines.append(f"# {text}")
                        elif 'heading 2' in style_name:
                            markdown_lines.append(f"## {text}")
                        elif 'heading 3' in style_name:
                            markdown_lines.append(f"### {text}")
                        elif 'heading 4' in style_name:
                            markdown_lines.append(f"#### {text}")
                        elif 'heading 5' in style_name:
                            markdown_lines.append(f"##### {text}")
                        elif 'heading 6' in style_name:
                            markdown_lines.append(f"###### {text}")
                        else:
                            markdown_lines.append(text)
                        
                        markdown_lines.append("")  # 添加空行
                        break
                        
            elif element.tag.endswith('tbl'):  # 表格
                # 找到对应的表格对象
                for table in doc.tables:
                    if table._element == element:
                        markdown_table = self._convert_table_to_markdown(table)
                        if markdown_table:
                            markdown_lines.extend(markdown_table)
                            markdown_lines.append("")  # 表格后添加空行
                        break
        
        return "\n".join(markdown_lines)
    
    def _convert_table_to_markdown(self, table) -> List[str]:
        """
        将Word表格转换为Markdown表格格式
        
        Args:
            table: python-docx表格对象
            
        Returns:
            Markdown表格行列表
        """
        if not table.rows:
            return []
            
        markdown_lines = []
        
        # 处理表格行
        for row_idx, row in enumerate(table.rows):
            row_cells = []
            for cell in row.cells:
                # 获取单元格文本，去除换行符和多余空格
                cell_text = cell.text.strip().replace('\n', ' ').replace('\r', ' ')
                # 转义Markdown特殊字符
                cell_text = cell_text.replace('|', '\\|')
                row_cells.append(cell_text)
            
            # 构建Markdown表格行
            markdown_row = '| ' + ' | '.join(row_cells) + ' |'
            markdown_lines.append(markdown_row)
            
            # 如果是第一行（表头），添加分隔符
            if row_idx == 0:
                separator_cells = ['---'] * len(row_cells)
                separator_row = '| ' + ' | '.join(separator_cells) + ' |'
                markdown_lines.append(separator_row)
        
        return markdown_lines

    def _convert_pdf_to_markdown(self, file_path: Path) -> str:
        """
        将PDF文档转换为Markdown
        
        Args:
            file_path: PDF文件路径
            
        Returns:
            Markdown内容
        """
        if pypandoc is not None:
            # 使用pypandoc进行转换（推荐方式）
            try:
                markdown_content = pypandoc.convert_file(
                    str(file_path), 
                    'markdown',
                    format='pdf',
                    extra_args=['--wrap=none']
                )
                return markdown_content
            except Exception as e:
                logger.warning(f"pypandoc PDF转换失败，使用备用方法: {str(e)}")
        
        # 备用方法：使用PyPDF2
        if PdfReader is not None:
            return self._pdf_to_markdown_fallback(file_path)
        
        raise RuntimeError("无可用的PDF转换方法")
    
    def _pdf_to_markdown_fallback(self, file_path: Path) -> str:
        """
        使用PyPDF2的备用转换方法
        
        Args:
            file_path: PDF文件路径
            
        Returns:
            Markdown内容
        """
        with open(file_path, 'rb') as file:
            pdf_reader = PdfReader(file)
            markdown_lines = []
            
            for page_num, page in enumerate(pdf_reader.pages, 1):
                text = page.extract_text()
                if text.strip():
                    # 添加页面标题
                    markdown_lines.append(f"## 第{page_num}页")
                    markdown_lines.append("")
                    
                    # 简单的文本处理
                    lines = text.split('\n')
                    for line in lines:
                        line = line.strip()
                        if line:
                            markdown_lines.append(line)
                    
                    markdown_lines.append("")  # 页面间添加空行
            
            return "\n".join(markdown_lines)
    
    def get_supported_formats(self) -> List[str]:
        """
        获取支持的文件格式
        
        Returns:
            支持的文件扩展名列表
        """
        return ['.docx', '.pdf']
    
    def is_supported(self, file_path: str) -> bool:
        """
        检查文件格式是否支持
        
        Args:
            file_path: 文件路径
            
        Returns:
            是否支持该格式
        """
        file_extension = Path(file_path).suffix.lower()
        return file_extension in self.get_supported_formats()
    
    def cleanup(self) -> None:
        """
        清理临时文件
        """
        try:
            import shutil
            if os.path.exists(self.temp_dir):
                shutil.rmtree(self.temp_dir)
        except Exception as e:
            logger.warning(f"清理临时文件失败: {str(e)}")
    
    def __del__(self):
        """
        析构函数，自动清理临时文件
        """
        self.cleanup()